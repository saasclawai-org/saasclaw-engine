"""Document extraction for wizard file uploads.

Extracts text content from PDF, Word, Excel, CSV, and text files
so the wizard agent can understand them without vision support.

Extracted text is injected into the user message as a labeled block,
keeping the original message and images intact.
"""

import base64
import csv
import io
import logging
import os
import tempfile

logger = logging.getLogger("studio.wizard_documents")

# Maximum characters to extract from any single document.
# Prevents context blowup from huge files.
MAX_EXTRACT_CHARS = 50_000

# Max file size for documents (25MB)
MAX_FILE_SIZE = 25 * 1024 * 1024

# Supported MIME types and their extensions
SUPPORTED_MIME_TYPES = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/msword": ".doc",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
    "application/vnd.ms-excel": ".xls",
    "text/csv": ".csv",
    "text/plain": ".txt",
    "text/markdown": ".md",
    "text/x-markdown": ".md",
}

# Extension-to-MIME fallback for when browser sends generic type
EXTENSION_MIME = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": "application/msword",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xls": "application/vnd.ms-excel",
    ".csv": "text/csv",
    ".txt": "text/plain",
    ".md": "text/markdown",
}


def is_document_mime(mime: str) -> bool:
    """Check if a MIME type is a supported document (not an image)."""
    return mime in SUPPORTED_MIME_TYPES


def is_document_file(filename: str) -> bool:
    """Check if a filename has a supported document extension."""
    _, ext = os.path.splitext(filename.lower())
    return ext in EXTENSION_MIME


def resolve_mime(mime: str, filename: str) -> str:
    """Resolve MIME type, falling back to extension-based detection."""
    if mime and mime != "application/octet-stream" and mime in SUPPORTED_MIME_TYPES:
        return mime
    _, ext = os.path.splitext(filename.lower())
    return EXTENSION_MIME.get(ext, mime)


def extract_text_from_document(data_b64: str, mime: str, filename: str) -> str | None:
    """Extract text content from a base64-encoded document.

    Args:
        data_b64: Base64-encoded file content.
        mime: MIME type (may be resolved from filename).
        filename: Original filename for extension fallback.

    Returns:
        Extracted text string, or None if extraction fails.
    """
    mime = resolve_mime(mime, filename)

    try:
        raw_bytes = base64.b64decode(data_b64)
    except Exception:
        logger.warning("Failed to decode base64 for document %s", filename)
        return None

    if len(raw_bytes) > MAX_FILE_SIZE:
        logger.warning("Document too large: %s (%d bytes)", filename, len(raw_bytes))
        return f"[Document '{filename}' is too large ({len(raw_bytes)} bytes). Max {MAX_FILE_SIZE} bytes.]"

    try:
        if mime == "application/pdf":
            text = _extract_pdf(raw_bytes, filename)
        elif mime in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            text = _extract_docx(raw_bytes, filename)
        elif mime in (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel",
        ):
            text = _extract_xlsx(raw_bytes, filename)
        elif mime == "text/csv":
            text = _extract_csv(raw_bytes, filename)
        elif mime in ("text/plain", "text/markdown", "text/x-markdown"):
            text = _extract_text(raw_bytes, filename)
        else:
            logger.warning("Unsupported document MIME: %s (%s)", mime, filename)
            return None

        if text and len(text) > MAX_EXTRACT_CHARS:
            text = text[:MAX_EXTRACT_CHARS] + f"\n\n[...truncated at {MAX_EXTRACT_CHARS} chars]"

        return text

    except Exception:
        logger.exception("Failed to extract text from %s", filename)
        return None


def _extract_pdf(raw_bytes: bytes, filename: str) -> str:
    """Extract text from a PDF using PyMuPDF (pymupdf)."""
    import pymupdf

    doc = pymupdf.Document(stream=raw_bytes, filetype="pdf")
    pages = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()
        if text.strip():
            pages.append(f"--- Page {page_num + 1} ---\n{text}")
    doc.close()

    if not pages:
        return f"[PDF '{filename}' contains no extractable text — it may be image-based.]"

    return f"[Contents of PDF: {filename}]\n\n" + "\n\n".join(pages)


def _extract_docx(raw_bytes: bytes, filename: str) -> str:
    """Extract text from a Word document (.docx/.doc)."""
    import docx

    doc = docx.Document(io.BytesIO(raw_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract tables
    tables_text = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            tables_text.append("\n".join(rows))

    parts = []
    if paragraphs:
        parts.append("\n\n".join(paragraphs))
    if tables_text:
        parts.append("[Tables]\n" + "\n\n".join(tables_text))

    if not parts:
        return f"[Word document '{filename}' contains no extractable text.]"

    return f"[Contents of Word document: {filename}]\n\n" + "\n\n".join(parts)


def _extract_xlsx(raw_bytes: bytes, filename: str) -> str:
    """Extract text from an Excel spreadsheet (.xlsx/.xls)."""
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(raw_bytes), read_only=True, data_only=True)
    sheets = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            # Skip completely empty rows
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            # Limit to 500 rows per sheet
            if len(rows) > 500:
                rows = rows[:500] + [f"... ({len(rows) - 500} more rows)"]
            sheets.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows))

    wb.close()

    if not sheets:
        return f"[Excel file '{filename}' contains no data.]"

    return f"[Contents of Excel: {filename}]\n\n" + "\n\n".join(sheets)


def _extract_csv(raw_bytes: bytes, filename: str) -> str:
    """Extract text from a CSV file."""
    content = raw_bytes.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(content))
    rows = []
    for i, row in enumerate(reader):
        if i >= 500:
            rows.append("... (more rows)")
            break
        rows.append(" | ".join(row))

    return f"[Contents of CSV: {filename}]\n\n" + "\n".join(rows)


def _extract_text(raw_bytes: bytes, filename: str) -> str:
    """Extract text from a plain text or markdown file."""
    content = raw_bytes.decode("utf-8", errors="replace")
    return f"[Contents of {filename}]\n\n{content}"


def format_document_text(extracted: str, filename: str) -> str:
    """Format extracted document text for injection into the user message.

    Returns a clearly labeled block that the LLM can understand as
    file content rather than user instructions.
    """
    return f"\n\n--- Attached Document: {filename} ---\n{extracted}\n--- End of {filename} ---\n"